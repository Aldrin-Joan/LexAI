import os
from docx import Document
from docx.shared import Pt
import re

def markdown_to_docx(md_path, docx_path):
    """
    Converts a simple markdown file to docx.
    Handles headers (# ## ###) and paragraphs.
    """
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by double newline to get blocks
    blocks = content.split('\n\n')

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # Handle Headers
        if block.startswith('# '):
            doc.add_heading(block[2:], level=1)
        elif block.startswith('## '):
            doc.add_heading(block[3:], level=2)
        elif block.startswith('### '):
            doc.add_heading(block[4:], level=3)
        elif block.startswith('#### '):
            doc.add_heading(block[5:], level=4)
        else:
            # Handle plain paragraphs
            # Very basic cleanup of markdown bold/italics (removal for docx style)
            # A more robust solution would use a library like pypandoc
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', block)
            clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)
            doc.add_paragraph(clean_text)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    src = r"c:\Users\ayush\Music\epics\final_submission\research_paper\FULL_RESEARCH_PAPER.md"
    dst = r"c:\Users\ayush\Music\epics\final_submission\research_paper\FULL_RESEARCH_PAPER.docx"
    markdown_to_docx(src, dst)
